import streamlit as st
import pandas as pd
#import numpy as np
import io
import samplics

from st_aggrid import AgGrid
from samplics.sampling import SampleSize
from samplics.utils.types import SizeMethod, PopParam
from fpdf import FPDF
#import sympy as sp
import matplotlib.pyplot as plt
import io
import os
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime

#from st_aggrid import AgGrid
#from st_aggrid.grid_options_builder import GridOptionsBuilder



# Vérifier si l'utilisateur a accès à cette page
st.title("Sondage stratifié")
if 'marge_strat' not in st.session_state:
    st.session_state.marge_strat = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse_strat' not in st.session_state:
    st.session_state.tauxReponse_strat = None
if 'tailleEffet_strat' not in st.session_state:
    st.session_state.tailleEffet_strat = None
if 'parametre_strat' not in st.session_state:
    st.session_state.parametre_strat = None
if 'methode_strat' not in st.session_state:
    st.session_state.methode_strat = None
if 'confiance_strat' not in st.session_state:
    st.session_state.confiance_strat = None
if 'sigma_strat' not in st.session_state:
    st.session_state.sigma_strat = None
if 'prop_est_strat' not in st.session_state:
    st.session_state.prop_est_strat = None
if 'calcul_strat' not in st.session_state:
    st.session_state.calcul_strat = False
if 'echantillon_strat' not in st.session_state:
    st.session_state.echantillon_strat = None
if 'base_echant_strat' not in st.session_state:
    st.session_state.base_echant_strat = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population_strat' not in st.session_state:
    st.session_state.population_strat = None
if 'variable_strat' not in st.session_state:
    st.session_state.variable_strat = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'modality_count' not in st.session_state:
    st.session_state.modality_count = None
if 'resultat' not in st.session_state:
    st.session_state.resultat = None
if 'tauxSondage_strat' not in st.session_state:
    st.session_state.tauxSondage_strat = None
if 'tot_taille' not in st.session_state:
    st.session_state.tot_taille = None
if 'nombre_tot' not in st.session_state:
    st.session_state.nombre_tot = None
if 'df_vide' not in st.session_state:
    st.session_state.df_vide = None
    
if 'marge' not in st.session_state:
    st.session_state.marge = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse' not in st.session_state:
    st.session_state.tauxReponse = None
if 'tailleEffet' not in st.session_state:
    st.session_state.tailleEffet = None
if 'parametre' not in st.session_state:
    st.session_state.parametre = None
if 'methode' not in st.session_state:
    st.session_state.methode = None
if 'confiance' not in st.session_state:
    st.session_state.confiance = None
if 'sigma' not in st.session_state:
    st.session_state.sigma = None
if 'prop_est' not in st.session_state:
    st.session_state.prop_est = None
if 'echantillon' not in st.session_state:
    st.session_state.echantillon = None
if 'calcul' not in st.session_state:
    st.session_state.calcul = False
if 'base_echant' not in st.session_state:
    st.session_state.base_echant = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population' not in st.session_state:
    st.session_state.population = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'type' not in st.session_state:
    st.session_state.type = None
if 'marge_strat' not in st.session_state:
    st.session_state.marge_strat = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse_strat' not in st.session_state:
    st.session_state.tauxReponse_strat = None
if 'tailleEffet_strat' not in st.session_state:
    st.session_state.tailleEffet_strat = None
if 'parametre_strat' not in st.session_state:
    st.session_state.parametre_strat = None
if 'methode_strat' not in st.session_state:
    st.session_state.methode_strat = None
if 'confiance_strat' not in st.session_state:
    st.session_state.confiance_strat = None
if 'sigma_strat' not in st.session_state:
    st.session_state.sigma_strat = None
if 'prop_est_strat' not in st.session_state:
    st.session_state.prop_est_strat = None
if 'calcul_strat' not in st.session_state:
    st.session_state.calcul_strat = False
if 'echantillon_strat' not in st.session_state:
    st.session_state.echantillon_strat = None
if 'base_echant_strat' not in st.session_state:
    st.session_state.base_echant_strat = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population_strat' not in st.session_state:
    st.session_state.population_strat = None
if 'variable_strat' not in st.session_state:
    st.session_state.variable_strat = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'modality_count' not in st.session_state:
    st.session_state.modality_count = None
if 'resultat' not in st.session_state:
    st.session_state.resultat = None
if 'tauxSondage' not in st.session_state:
    st.session_state.tauxSondage = None
if 'tauxSondage_strat' not in st.session_state:
    st.session_state.tauxSondage_strat = None
if 'tot_taille' not in st.session_state:
    st.session_state.tot_taille = None

st.info("Les champs avec l'astérisque (*) rouge doivent être obligatoirement remplis")
st.markdown("##### Paramètre à estimer <span style='color:red;'>*</span>", unsafe_allow_html=True)
st.session_state.parametre_strat = st.selectbox('Paramètre à estimer', [None, 'Proportion', 'Moyenne'], label_visibility='collapsed')

cola, colb = st.columns(2)
with cola:
    st.markdown("##### Type de sondage stratifié <span style='color:red;'>*</span>", unsafe_allow_html=True)
    st.session_state.type_alloc = st.radio("Veuillez choisir le type de sondage stratifié souhaité", options = [None, 'Allocation proportionnelle', 'Calcul par strate'], label_visibility='collapsed')

if st.session_state.parametre_strat is not None and st.session_state.type_alloc is not None:
    
    with colb:
        if st.session_state.df is None:
            var = st.text_input("Quelle est la variable de stratification ?")
            n= int(st.number_input(f"Nombre de modalités (strates) dans {var} ?"))
            noms_colonnes = ["strate", "Taille"]
            # Créer un dataframe vide
            st.session_state.df_vide = pd.DataFrame([[None] * 2 for _ in range(n)], columns=noms_colonnes)

        if st.session_state.df is not None:
            # Filtrer les variables catégorielles (type 'object' ou avec moins de 6 valeurs uniques)
            variable_categ = [col for col in st.session_state.df.columns if st.session_state.df[col].dtype == 'object' or st.session_state.df[col].nunique() < 6]
            st.markdown("##### Variable de stratification (catégorielle) <span style='color:red;'>*</span>", unsafe_allow_html=True)
            # Vérifier s'il y a des variables catégorielles
            if len(variable_categ) > 0:
                st.session_state.variable_strat = st.selectbox('**Choisissez la variable de stratification (catégorielle)**', variable_categ, label_visibility='collapsed')
                if st.session_state.variable_strat is not None:
                    st.session_state.modality_count = pd.DataFrame(st.session_state.df[st.session_state.variable_strat].value_counts().reset_index())
                    st.session_state.modality_count.columns = [st.session_state.variable_strat, 'Taille']
                    # Afficher le tableau de données si il s'agit d'une allocation proportionnelle
                    #if st.session_state.type_alloc == 'Allocation proportionnelle':
                    #    st.dataframe(st.session_state.modality_count)

    if st.session_state.df is None:
        if st.session_state.type_alloc == "Allocation proportionnelle":
            st.write(f"Veuillez remplir les modalités de la variable {var} dans strates et donner la population dans chaque strate :")
            st.session_state.modality_count = st.data_editor(st.session_state.df_vide)
            st.session_state.modality_count['Taille'] = pd.to_numeric(st.session_state.modality_count['Taille'], errors='coerce')
            # Calculer la taille totale de la population pour l'allocation proportionnelle plus tard
            st.session_state.nombre_tot = st.session_state.modality_count['Taille'].sum()



    if st.session_state.type_alloc == 'Calcul par strate':
        if st.session_state.df is None:
            st.session_state.df_vide['Effet de plan'] = 0.0
            if st.session_state.parametre_strat == "Proportion":
                st.session_state.df_vide['Proportion estimée'] = 0.0
            if st.session_state.parametre_strat == "Moyenne":
                st.session_state.df_vide['Ecart-type estimé'] = 0.0
        elif st.session_state.df is not None:
            st.session_state.modality_count['Effet de plan'] = 0.0
            if st.session_state.parametre_strat == "Proportion":
                st.session_state.modality_count['Proportion estimée'] = 0.0
            if st.session_state.parametre_strat == "Moyenne":
                st.session_state.modality_count['Ecart-type estimé'] = 0.0


        st.write("**Remplissez les données manquantes du tableau**  <span style='color:red;'>*</span>", unsafe_allow_html=True)
        with st.expander("En savoir plus ..."):
            st.info("Si vous disposez d'études similaires précédentes vous pouvez utiliser ces résultats pour la proportion estimée. Sinon la valeur prise par défaut est 0.5.")
            st.info("\nL'effet de plan est un facteur utilisé dans les enquêtes pour ajuster la taille d'échantillon nécessaire lorsque le plan d'échantillonnage diffère d'un simple échantillonnage aléatoire. Il mesure l'impact du plan d'échantillonnage sur la précision des estimations. Si vous utilisez un plan d'échantillonnage stratifié ou en grappes, par exemple, l'effet de plan prend en compte la corrélation au sein des groupes ou des strates, ce qui peut augmenter la variance des estimations. Un effet de plan supérieur à 1 signifie que la taille de l'échantillon doit être augmentée pour maintenir la même précision. Comme nous sommes dans le cas d'un sondage aléatoire simple, l'effet de plan est de 1")
            st.info(" \n La valeur par défaut de la taille est le nombre d'individus par strate de la base. ")
        if st.session_state.df is None:
            st.session_state.modality_count = st.data_editor(st.session_state.df_vide)
            st.session_state.modality_count['Taille'] = pd.to_numeric(st.session_state.modality_count['Taille'], errors='coerce')
        elif st.session_state.df is not None:
            st.session_state.modality_count = st.data_editor(st.session_state.modality_count)
            
        # Calculer la taille totale de la population pour le calcul par strate
        st.session_state.nombre_tot = st.session_state.modality_count['Taille'].sum()
        
        # Récupération des valeurs du tableau pour les calculs
        if st.session_state.parametre_strat == "Moyenne":
            if st.session_state.df is not None:
                st.session_state.sigma_strat = st.session_state.modality_count[[st.session_state.variable_strat, 'Ecart-type estimé']].to_dict()
            elif st.session_state.df is None:
                st.session_state.sigma_strat = st.session_state.modality_count[['strate', 'Ecart-type estimé']].to_dict()
            st.session_state.sigma_strat_nombre = st.session_state.sigma_strat['Ecart-type estimé']

        #Extraction de la taille de la population pour chaque strate
        if st.session_state.df is not None:
            st.session_state.population_strat = st.session_state.modality_count[[st.session_state.variable_strat, 'Taille']].to_dict()
        elif st.session_state.df is  None:
            st.session_state.population_strat = st.session_state.modality_count[['strate', 'Taille']].to_dict()
        st.session_state.population_strat_nombre = st.session_state.population_strat['Taille']
        
        #Extraction de l'effet de plan pour chaque strate
        if st.session_state.df is not None:
            st.session_state.tailleEffet_strat = st.session_state.modality_count[[st.session_state.variable_strat, 'Effet de plan']].to_dict()
        elif st.session_state.df is None:
            st.session_state.tailleEffet_strat = st.session_state.modality_count[['strate', 'Effet de plan']].to_dict()
        st.session_state.tailleEffet_strat_nombre = st.session_state.tailleEffet_strat['Effet de plan']
        if st.session_state.df is not None:
            st.session_state.tailleEffet_strat_label = st.session_state.tailleEffet_strat[st.session_state.variable_strat]
        elif st.session_state.df is None:
            st.session_state.tailleEffet_strat_label = st.session_state.tailleEffet_strat['strate']

        if st.session_state.parametre_strat == "Proportion":
            if st.session_state.df is not None:
                st.session_state.prop_est_strat = st.session_state.modality_count[[st.session_state.variable_strat, 'Proportion estimée']].to_dict()
            elif st.session_state.df is None:
                st.session_state.prop_est_strat = st.session_state.modality_count[['strate', 'Proportion estimée']].to_dict()
            st.session_state.prop_est_strat_nombre = st.session_state.prop_est_strat['Proportion estimée']

        colc, cold, cole = st.columns(3)
        with colc:
            if st.session_state.parametre_strat == "Proportion":
                st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
                st.session_state.marge_strat = st.number_input("Marge d'erreur",min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')


            if st.session_state.parametre_strat == "Moyenne":
                st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
                st.session_state.marge_strat = st.number_input("Marge d'erreur moyenne ", min_value= 0, label_visibility='collapsed')

        with cold:
            st.markdown("##### Niveau de confiance <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Le niveau de confiance représente la probabilité que vous êtes prêt à accepter quant à la possibilité que votre estimation soit incorrecte.")
            st.session_state.confiance_strat = st.number_input("Niveau de confiance", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')

        with cole:
            st.markdown("##### Taux de réponse <span style='color:red;'>*</span> ", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Le taux de réponse est le pourcentage de personnes ayant répondu à un sondage par rapport au nombre total de personnes sollicitées. Assurez-vous d'atteindre ce taux de réponse au risque de ne pas avoir des estimations précises.")
            st.session_state.tauxReponse_strat = st.number_input("Taux de réponse", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", value=0.90, label_visibility='collapsed')


        # Calcul de la taille dans le cas d'une proportion
        if st.session_state.parametre_strat == "Proportion":
            # Initialisation de la méthode
            st.session_state.methode_strat = SampleSize(param=PopParam.prop, method=SizeMethod.wald, strat=True)
            # Vérification des conditions avant le calcul
            if (
                st.session_state.methode_strat is not None and
                all(valeur is not None for valeur in st.session_state.tailleEffet_strat_nombre.values()) and
                all(valeur is not None for valeur in st.session_state.prop_est_strat_nombre.values()) and
                all(valeur is not None for valeur in st.session_state.population_strat_nombre.values()) and
                st.session_state.tauxReponse_strat != 0 and 
                st.session_state.marge_strat != 0
                ):
                # Fonction pour calculer la taille d'échantillon
                def calcul_taille():
                    st.session_state.methode_strat.calculate(
                        target=st.session_state.prop_est_strat_nombre, 
                        half_ci=st.session_state.marge_strat,
                        deff=st.session_state.tailleEffet_strat_nombre, 
                        resp_rate=st.session_state.tauxReponse_strat,
                        pop_size=st.session_state.population_strat_nombre,
                        alpha=1 - st.session_state.confiance_strat
                    )
                    st.session_state.echantillon_strat = st.session_state.methode_strat.samp_size
                # Bouton pour lancer le calcul
                st.session_state.calcul_strat = st.button(
                    "Calculer la taille", 
                    help="Cliquer pour lancer le calcul de la taille", 
                    on_click=calcul_taille
                )

        # Calcul de la taille dans le cas d'une Moyenne
        if st.session_state.parametre_strat == "Moyenne":
            # Initialisation de la méthode
            st.session_state.methode_strat = SampleSize(param=PopParam.mean, method=SizeMethod.wald, strat=True)
            
            # Vérification des conditions avant le calcul
            if (
                st.session_state.methode_strat is not None and
                all(valeur is not None for valeur in st.session_state.tailleEffet_strat_nombre.values()) and
                all(valeur is not None for valeur in st.session_state.sigma_strat_nombre.values()) and
                all(valeur is not None for valeur in st.session_state.population_strat_nombre.values()) and
                st.session_state.tauxReponse_strat != 0 and 
                st.session_state.marge_strat != 0
                ):
                # Fonction pour calculer la taille d'échantillon
                def calcul_taille():
                    st.session_state.methode_strat.calculate(
                        target=st.session_state.sigma_strat_nombre,
                        half_ci=st.session_state.marge_strat,
                        deff=st.session_state.tailleEffet_strat_nombre,
                        resp_rate=st.session_state.tauxReponse_strat,
                        pop_size=st.session_state.population_strat_nombre,
                        alpha=1 - st.session_state.confiance_strat
                    )
                    st.session_state.echantillon_strat = st.session_state.methode_strat.samp_size
                # Bouton pour lancer le calcul
                st.session_state.calcul_strat = st.button(
                    "Calculer la taille", 
                    help="Cliquer pour lancer le calcul de la taille", 
                    on_click=calcul_taille
                )

        
        if st.session_state.calcul_strat:
            st.markdown("<h1 style='color:#2fd5cf; text-align: center;'>Résultats </h1>", unsafe_allow_html=True)
            colj, colk, coll = st.columns(3)
            with colj:
                # Fonction pour combiner les valeurs
                def combine_valeurs(key):
                    return (st.session_state.tailleEffet_strat_label[key], st.session_state.echantillon_strat[key])

                # Obtenir toutes les clés
                cle = st.session_state.tailleEffet_strat_label.keys()
                # Utiliser map pour appliquer la fonction à chaque clé
                combined_values = map(combine_valeurs, cle)
                # Convertir le résultat de map en liste
                combined_list = list(combined_values)
                # Créer un DataFrame
                st.session_state.resultat = pd.DataFrame(combined_list, index=cle, columns=['Strates', "Taille de l'échantillon"])
                st.session_state.resultat['Taux de sondage'] = round(st.session_state.resultat["Taille de l'échantillon"] / st.session_state.modality_count["Taille"], 5)
                
                tot_echant = st.session_state.resultat["Taille de l'échantillon"].sum()
                tot_sond = round(tot_echant/st.session_state.nombre_tot,4)
                st.session_state.resultat.loc['Total'] = ['Total',tot_echant,tot_sond]
                st.dataframe(st.session_state.resultat)

                if st.session_state.df is not None:
                    with st.expander('Afficher plus'):
                        # Effectuer les tirages dans chaque strate
                        sampled_dataframes = []
                        for index, row in st.session_state.resultat.iterrows():
                            modality = row["Strates"]
                            sample_size = row["Taille de l'échantillon"]
                            
                            # Filtrer la base par strate/modality
                            if st.session_state.df is not None:
                                subset_strata = st.session_state.df[st.session_state.df[st.session_state.variable_strat] == modality]
                            elif st.session_state.df is None:
                                subset_strata = st.session_state.df[st.session_state.df['strate'] == modality]
                            # Tirer un échantillon dans chaque strate
                            if len(subset_strata) >= sample_size:  # Vérifier que la taille de la strate permet le tirage
                                sampled_strata = subset_strata.sample(n=sample_size, random_state=42)
                                sampled_dataframes.append(sampled_strata)
                            
                        # Combiner tous les échantillons tirés en un seul DataFrame
                        final_sample = pd.concat(sampled_dataframes)

                        # Stocker dans session_state pour utilisation ultérieure
                        st.session_state.base_echant_strat = final_sample

                        st.write("**Échantillon stratifié tiré** :")
                        AgGrid(st.session_state.base_echant_strat)

            with colk:
                def generate_latex_image(formula, fontsize=30):
                    fig, ax = plt.subplots()
                    ax.text(0.5, 0.5, f'${formula}$', fontsize=fontsize, ha='center', va='center')
                    ax.axis('off')

                    # Sauvegarder l'image dans un buffer mémoire
                    img_buffer = io.BytesIO()
                    plt.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
                    plt.close(fig)
                    img_buffer.seek(0)
                    return img_buffer

                def generate_stratified_report():
                    # Créer un document Word
                    doc = Document()
                    # Ajouter un titre
                    doc.add_heading('Rapport des résultats (Sondage Stratifié)', 0)
                    
                    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    doc.add_paragraph(f'Généré le : {current_time}')   
                    doc.add_paragraph("L'application SamPower est un outil puissant dédié à la réalisation de sondages et à l'estimation des paramètres statistiques essentiels. Conçue pour être à la fois intuitive et accessible, elle permet aux utilisateurs d'effectuer des échantillonnages de manière rapide et efficace, que ce soit pour des sondages aléatoires simples ou stratifiés. Son importance réside dans sa capacité à offrir des résultats fiables tout en simplifiant le processus complexe de calcul de taille d'échantillon, un élément clé pour garantir des résultats statistiquement représentatifs dans le cadre d'études quantitatives. Le présent rapport est structuré de la manière suivante :")
                    doc.add_paragraph("- Présentation des paramètres d'entrée et des résultats : Cette section regroupe l'ensemble des résultats obtenus ainsi que les différents paramètres entrés par l'utilisateur.")            
                    doc.add_paragraph("- Présentation des explications et de la méthodologie utilisée : Ici, nous détaillons la méthodologie employée pour générer les résultats, notamment la logique derrière les formules et les méthodes statistiques appliquées.")
                    
                    # Ajouter des informations sur la méthode
                    doc.add_heading("Présentation des paramètres d'entrée et des résultats", level = 1)
                    doc.add_heading('Paramètres à estimer :', level=2)
                    doc.add_paragraph(f"Paramètre : {st.session_state.parametre_strat}")

                    # Ajouter les paramètres d'entrée
                    doc.add_heading('Paramètres d\'entrée :', level=2)
                    doc.add_paragraph(f"Niveau de confiance : {st.session_state.confiance_strat * 100}%")
                    doc.add_paragraph(f"Erreur Marginale : {st.session_state.marge_strat}")
                    doc.add_paragraph(f"Taux de réponse : {st.session_state.tauxReponse_strat * 100}%")
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Taux de sondage : {st.session_state.tauxSondage_strat * 100}%")

                
                    if st.session_state.type_alloc == 'Calcul par strate':
                        if isinstance(st.session_state.modality_count, pd.DataFrame):
                            # Ajouter les paramètres sous forme de tableau dans le document Word
                            doc.add_heading('Taille et effet de plan :', level=3)
                            table = doc.add_table(rows=1, cols=len(st.session_state.modality_count.columns))
                            hdr_cells = table.rows[0].cells
                            for i, column in enumerate(st.session_state.modality_count.columns):
                                hdr_cells[i].text = column

                            for index, row in st.session_state.modality_count.iterrows():
                                row_cells = table.add_row().cells
                                for i, value in enumerate(row):
                                    row_cells[i].text = str(value)


                    # Afficher les résultats d'entrée
                    doc.add_heading('Résultats :', level=2)
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Taille de l'échantillon calculée : {st.session_state.echantillon_strat}")
                    if st.session_state.type_alloc == 'Calcul par strate':
                        result_df = st.session_state.resultat
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        result_df = st.session_state.modality_count
                    doc.add_page_break()
                    if isinstance(result_df, pd.DataFrame):
                        # Ajouter les paramètres sous forme de tableau dans le document Word
                        doc.add_heading('Tableau des résultats', level=3)
                        table = doc.add_table(rows=1, cols=len(result_df.columns))
                        hdr_cells = table.rows[0].cells
                        for i, column in enumerate(result_df.columns):
                            hdr_cells[i].text = column

                        for index, row in result_df.iterrows():
                            row_cells = table.add_row().cells
                            for i, value in enumerate(row):
                                row_cells[i].text = str(value)
                    # Ajouter des explications
                    
                    doc.add_heading('Explications et méthodologie utilisée', level=1)
                    doc.add_heading('Explication des résultats :', level=2)
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Suite à l'application de la méthode de sondage stratifié, l'échantillon total a été fixé à {st.session_state.echantillon_strat}. Cet échantillon a été réparti de manière proportionnelle entre les différentes strates en fonction de leur taille dans la population de référence comme présenté dans le tableau des résultats.Cette répartition garantit que chaque strate est représentée de manière adéquate dans l'échantillon final, permettant ainsi des estimations précises et fiables pour chacune des strates tout en minimisant les biais potentiels liés à la taille des sous-populations. En conséquence, les résultats obtenus à partir de cet échantillon seront statistiquement représentatifs de l'ensemble de la population cible tout en respectant la structure spécifique des strates.")            
                    if st.session_state.type_alloc == 'Calcul par strate':
                        doc.add_paragraph(f"Dans le cadre de ce sondage stratifié, la taille de l'échantillon total a été directement déterminée au sein de chaque strate selon des critères spécifiques à chacune d'elles, tels que leur proportion ou leur variabilité interne. L'échantillon est donc de {st.session_state.tot_taille}. Cette méthode permet de s'assurer que chaque strate est représentée de manière proportionnée à ses caractéristiques propres, garantissant ainsi une précision accrue des estimations dans chaque sous-groupe de la population. En optant pour un calcul direct des tailles d'échantillon dans chaque strate, nous maximisons la représentativité de l'échantillon pour chaque groupe cible, tout en prenant en compte les particularités de chaque sous-population. Ce processus offre des résultats plus détaillés et pertinents pour chaque strate, assurant ainsi une meilleure compréhension des dynamiques internes de la population étudiée. ")
                    
                    # Formule de la méthode de Wald
                    if st.session_state.parametre_strat == 'Proportion':
                        doc.add_heading('Formule de calcul',2)
                        doc.add_paragraph('La formule de calcul utilisée est la suivante:')
                        formula = r'n = \frac{Z^2 \cdot p \cdot (1 - p)}{e^2}'
                        img_buffer = generate_latex_image(formula)
                        doc.add_picture(img_buffer, width=Inches(3))
                        doc.add_paragraph("Avec n la taille de l'échantillon, Z le niveau de confiance, p la proportion estimée et e la marge d'erreur")

                        doc.add_heading('Méthodologie', 2)
                        doc.add_paragraph("La méthode utilisée dans le calcul de la taille de l'échantillon est la méthode de Cochran")

                    # doc.add_picture('path_to_graph.png', width=Inches(5))
                    # Sauvegarder le fichier Word en mémoire
                    report_io = io.BytesIO()
                    doc.save(report_io)
                    report_io.seek(0)
                    return report_io

                # Télécharger le rapport si l'utilisateur clique sur le bouton
                if st.session_state.type == 'Stratifie':
                    report_file = generate_stratified_report()
                # Télécharger le rapport
                st.download_button(label="Télécharger le Rapport", data=report_file, file_name="rapport_sondage.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            with coll:
                def to_excel(df):
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                    return output.getvalue()

                if st.session_state.base_echant_strat is not None:
                    if st.session_state.type == 'Stratifie':
                        excel_data = to_excel(st.session_state.base_echant_strat)

                # Ajouter un bouton de téléchargement
                    st.download_button(
                        label="Télécharger l'échantillon sous format Excel",
                        data=excel_data,
                        file_name='Echantillon.xlsx',
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    )


    if st.session_state.type_alloc == 'Allocation proportionnelle':
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.session_state.parametre_strat == "Moyenne":
                st.markdown("##### Ecart-type estimé <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("Si vous avez accès à des données provenant d'études antérieures sur la même population ou des populations similaires, vous pouvez utiliser l'écart-type observé dans ces études comme estimation de 𝜎σ.")
                st.session_state.sigma_strat = st.number_input("Marge d'erreur ", min_value = 0, label_visibility='collapsed')

            if st.session_state.parametre_strat == "Proportion":
                st.markdown("##### Proportion estimée <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("Si vous disposez d'études similaires précédentes vous pouvez utiliser ces résultats. Sinon la valeur prise par défaut est 0.5")
                st.session_state.prop_est_strat = st.number_input("Proportion estimée", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", value=0.5, label_visibility='collapsed')

            if st.session_state.parametre_strat == "Proportion":
                st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
                st.session_state.marge_strat = st.number_input("Marge d'erreur", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", value=0.0005, label_visibility='collapsed')

            if st.session_state.parametre_strat == "Moyenne":
                st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
                with st.expander("En savoir plus ..."):
                    st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
                st.session_state.marge_strat = st.number_input("Marge d'erreur moy ", min_value= 0, label_visibility='collapsed')

        with col2:
            st.markdown("##### Niveau de confiance <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Le niveau de confiance représente la probabilité que vous êtes prêt à accepter quant à la possibilité que votre estimation soit incorrecte.")
            st.session_state.confiance_strat = st.number_input("Niveau de confiance",min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", value=0.95, label_visibility='collapsed')

            st.markdown("##### Taille de la population <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("La valeur par défaut de la taille est le nombre d'individus de la base s'il y en a")
            if st.session_state.df is not None:
                st.session_state.population_strat = st.number_input("Taille de la population", min_value = 1, value=st.session_state.taillePop, label_visibility='collapsed')
            if st.session_state.df is None and st.session_state.nombre_tot > 1:
                st.session_state.population_strat = st.number_input("Taille de la population", min_value = 1, value=st.session_state.nombre_tot, label_visibility='collapsed')
        
        with col3:
            st.markdown("##### Taux de réponse <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Le taux de réponse est le pourcentage de personnes ayant répondu à un sondage par rapport au nombre total de personnes sollicitées. Assurez-vous d'atteindre ce taux de réponse au risque de ne pas avoir des estimations précises.")
            st.session_state.tauxReponse_strat = st.number_input("Taux de réponse", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", value=0.90, label_visibility='collapsed')

            st.markdown("##### Effet de plan <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("L'effet de plan est un facteur utilisé dans les enquêtes pour ajuster la taille d'échantillon nécessaire lorsque le plan d'échantillonnage diffère d'un simple échantillonnage aléatoire. Il mesure l'impact du plan d'échantillonnage sur la précision des estimations. Si vous utilisez un plan d'échantillonnage stratifié ou en grappes, par exemple, l'effet de plan prend en compte la corrélation au sein des groupes ou des strates, ce qui peut augmenter la variance des estimations. Un effet de plan supérieur à 1 signifie que la taille de l'échantillon doit être augmentée pour maintenir la même précision. Comme nous sommes dans le cas d'un sondage aléatoire simple, l'effet de plan est de 1")
            st.session_state.tailleEffet_strat =  st.selectbox("Effet de plan",[0.5,0.75,0.90,1],label_visibility='collapsed')



        # Calcul de la taille dans le cas d'une proportion
        if st.session_state.parametre_strat == "Proportion":
            st.session_state.methode_strat = SampleSize(param=PopParam.prop, method=SizeMethod.wald, strat=False)
            if st.session_state.methode_strat is not None and st.session_state.tailleEffet_strat != 0 and st.session_state.tauxReponse_strat != 0 and st.session_state.population_strat != 0 and st.session_state.marge_strat != 0:
                def calcul_taille():
                    st.session_state.methode_strat.calculate(
                        target = st.session_state.prop_est_strat, 
                        half_ci=st.session_state.marge_strat,
                        deff=st.session_state.tailleEffet_strat, 
                        resp_rate=st.session_state.tauxReponse_strat,
                        pop_size = st.session_state.population_strat,
                        alpha = 1 - st.session_state.confiance_strat)
                    st.session_state.echantillon_strat = st.session_state.methode_strat.samp_size
                st.session_state.calcul_strat = st.button("Calculer la taille",help="Cliquer pour lancer le calcul de la taille",on_click=calcul_taille)


        # Calcul de la taille dans le cas d'une moyenne
        if st.session_state.parametre_strat == "Moyenne":
            st.session_state.methode_strat = SampleSize(param=PopParam.mean, method=SizeMethod.wald, strat=False)
            if st.session_state.methode_strat is not None and st.session_state.tailleEffet_strat != 0 and st.session_state.tauxReponse_strat != 0 and st.session_state.taillePop_strat != 0 and st.session_state.marge_strat != 0:
                def calcul_taille():
                    st.session_state.methode_strat.calculate(
                        sigma = st.session_state.sigma_strat,
                        half_ci=st.session_state.marge_strat,
                        deff=st.session_state.tailleEffet_strat, 
                        resp_rate=st.session_state.tauxReponse_strat,
                        pop_size = st.session_state.population_strat,
                        alpha = 1 - st.session_state.confiance_strat)
                    st.session_state.echantillon_strat = st.session_state.methode_strat.samp_size
                st.session_state.calcul_strat = st.button("Calculer la taille",help="Cliquer pour lancer le calcul de la taille",on_click=calcul_taille)
        
        if st.session_state.calcul_strat:
            st.markdown("<h1 style='color:#2fd5cf; text-align: center;'>Résultats </h1>", unsafe_allow_html=True)
            colm, coln, colo = st.columns(3)
            with colm:
                st.write("**Taille de l'échantillon**: ",st.session_state.echantillon_strat)
                st.session_state.tauxSondage_strat = round(st.session_state.echantillon_strat / st.session_state.population_strat,4)
                st.write("Le taux de sondage est: ",st.session_state.tauxSondage_strat)
                st.session_state.modality_count["Taille de l'echantillon"] = round(st.session_state.modality_count['Taille'] * st.session_state.tauxSondage_strat).astype(int)
                st.session_state.modality_count['Taux de sondage'] = round(st.session_state.modality_count["Taille de l'echantillon"] / st.session_state.modality_count['Taille'],4)
                tot_nomb = st.session_state.modality_count['Taille'].sum()
                st.session_state.tot_taille = st.session_state.modality_count["Taille de l'echantillon"].sum()
                tot_tx = st.session_state.tot_taille / tot_nomb
                st.session_state.modality_count.loc['Total'] = ['Total',tot_nomb,st.session_state.tot_taille,tot_tx]
                st.dataframe(st.session_state.modality_count)

                if st.session_state.df is not None:
                    with st.expander('Afficher plus'):
                    # Effectuer les tirages dans chaque strate
                        sampled_dataframes = []
                        for index, row in st.session_state.modality_count.iterrows():
                            if st.session_state.df is not None:
                                modality = row[st.session_state.variable_strat]
                            elif st.session_state.df is None:
                                modality = row['strate']
                            sample_size = row["Taille de l'echantillon"]
                            
                            # Filtrer la base par strate/modalité
                            if st.session_state.df is not None:
                                subset_strata = st.session_state.df[st.session_state.df[st.session_state.variable_strat] == modality]
                            elif st.session_state.df is None:
                                subset_strata = st.session_state.df[st.session_state.df['strate'] == modality]
                        
                            # Tirer un échantillon dans chaque strate
                            if len(subset_strata) >= sample_size:  # Vérifier que la taille de la strate permet le tirage
                                sampled_strata = subset_strata.sample(n=sample_size, random_state=42)
                                sampled_dataframes.append(sampled_strata)

                        # Combiner tous les échantillons tirés en un seul DataFrame
                        final_sample = pd.concat(sampled_dataframes)

                        # Stocker dans session_state pour utilisation ultérieure
                        st.session_state.base_echant_strat = final_sample

                        st.write("Échantillon stratifié tiré :")
                        AgGrid(st.session_state.base_echant_strat)
            
            with coln:
                def generate_latex_image(formula, fontsize=30):
                    fig, ax = plt.subplots()
                    ax.text(0.5, 0.5, f'${formula}$', fontsize=fontsize, ha='center', va='center')
                    ax.axis('off')

                    # Sauvegarder l'image dans un buffer mémoire
                    img_buffer = io.BytesIO()
                    plt.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
                    plt.close(fig)
                    img_buffer.seek(0)
                    return img_buffer

                def generate_stratified_report():
                    # Créer un document Word
                    doc = Document()
                    # Ajouter un titre
                    doc.add_heading('Rapport des résultats (Sondage Stratifié)', 0)
                    
                    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    doc.add_paragraph(f'Généré le : {current_time}')   
                    doc.add_paragraph("L'application SamPower est un outil puissant dédié à la réalisation de sondages et à l'estimation des paramètres statistiques essentiels. Conçue pour être à la fois intuitive et accessible, elle permet aux utilisateurs d'effectuer des échantillonnages de manière rapide et efficace, que ce soit pour des sondages aléatoires simples ou stratifiés. Son importance réside dans sa capacité à offrir des résultats fiables tout en simplifiant le processus complexe de calcul de taille d'échantillon, un élément clé pour garantir des résultats statistiquement représentatifs dans le cadre d'études quantitatives. Le présent rapport est structuré de la manière suivante :")
                    doc.add_paragraph("- Présentation des paramètres d'entrée et des résultats : Cette section regroupe l'ensemble des résultats obtenus ainsi que les différents paramètres entrés par l'utilisateur.")            
                    doc.add_paragraph("- Présentation des explications et de la méthodologie utilisée : Ici, nous détaillons la méthodologie employée pour générer les résultats, notamment la logique derrière les formules et les méthodes statistiques appliquées.")
                    
                    # Ajouter des informations sur la méthode
                    doc.add_heading("Présentation des paramètres d'entrée et des résultats", level = 1)
                    doc.add_heading('Paramètres à estimer :', level=2)
                    doc.add_paragraph(f"Paramètre : {st.session_state.parametre_strat}")

                    # Ajouter les paramètres d'entrée
                    doc.add_heading('Paramètres d\'entrée :', level=2)
                    doc.add_paragraph(f"Niveau de confiance : {st.session_state.confiance_strat * 100}%")
                    doc.add_paragraph(f"Erreur Marginale : {st.session_state.marge_strat}")
                    doc.add_paragraph(f"Taux de réponse : {st.session_state.tauxReponse_strat * 100}%")
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Taux de sondage : {st.session_state.tauxSondage_strat * 100}%")

                
                    if st.session_state.type_alloc == 'Calcul par strate':
                        if isinstance(st.session_state.modality_count, pd.DataFrame):
                            # Ajouter les paramètres sous forme de tableau dans le document Word
                            doc.add_heading('Taille et effet de plan :', level=3)
                            table = doc.add_table(rows=1, cols=len(st.session_state.modality_count.columns))
                            hdr_cells = table.rows[0].cells
                            for i, column in enumerate(st.session_state.modality_count.columns):
                                hdr_cells[i].text = column

                            for index, row in st.session_state.modality_count.iterrows():
                                row_cells = table.add_row().cells
                                for i, value in enumerate(row):
                                    row_cells[i].text = str(value)


                    # Afficher les résultats d'entrée
                    doc.add_heading('Résultats :', level=2)
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Taille de l'échantillon calculée : {st.session_state.echantillon_strat}")
                    if st.session_state.type_alloc == 'Calcul par strate':
                        result_df = st.session_state.resultat
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        result_df = st.session_state.modality_count
                    doc.add_page_break()
                    if isinstance(result_df, pd.DataFrame):
                        # Ajouter les paramètres sous forme de tableau dans le document Word
                        doc.add_heading('Tableau des résultats', level=3)
                        table = doc.add_table(rows=1, cols=len(result_df.columns))
                        hdr_cells = table.rows[0].cells
                        for i, column in enumerate(result_df.columns):
                            hdr_cells[i].text = column

                        for index, row in result_df.iterrows():
                            row_cells = table.add_row().cells
                            for i, value in enumerate(row):
                                row_cells[i].text = str(value)
                    # Ajouter des explications
                    
                    doc.add_heading('Explications et méthodologie utilisée', level=1)
                    doc.add_heading('Explication des résultats :', level=2)
                    if st.session_state.type_alloc == 'Allocation proportionnelle':
                        doc.add_paragraph(f"Suite à l'application de la méthode de sondage stratifié, l'échantillon total a été fixé à {st.session_state.echantillon_strat}. Cet échantillon a été réparti de manière proportionnelle entre les différentes strates en fonction de leur taille dans la population de référence comme présenté dans le tableau des résultats.Cette répartition garantit que chaque strate est représentée de manière adéquate dans l'échantillon final, permettant ainsi des estimations précises et fiables pour chacune des strates tout en minimisant les biais potentiels liés à la taille des sous-populations. En conséquence, les résultats obtenus à partir de cet échantillon seront statistiquement représentatifs de l'ensemble de la population cible tout en respectant la structure spécifique des strates.")            
                    if st.session_state.type_alloc == 'Calcul par strate':
                        doc.add_paragraph(f"Dans le cadre de ce sondage stratifié, la taille de l'échantillon total a été directement déterminée au sein de chaque strate selon des critères spécifiques à chacune d'elles, tels que leur proportion ou leur variabilité interne. L'échantillon est donc de {st.session_state.tot_taille}. Cette méthode permet de s'assurer que chaque strate est représentée de manière proportionnée à ses caractéristiques propres, garantissant ainsi une précision accrue des estimations dans chaque sous-groupe de la population. En optant pour un calcul direct des tailles d'échantillon dans chaque strate, nous maximisons la représentativité de l'échantillon pour chaque groupe cible, tout en prenant en compte les particularités de chaque sous-population. Ce processus offre des résultats plus détaillés et pertinents pour chaque strate, assurant ainsi une meilleure compréhension des dynamiques internes de la population étudiée. ")
                    
                    # Formule de la méthode de Wald
                    if st.session_state.parametre_strat == 'Proportion':
                        doc.add_heading('Formule de calcul',2)
                        doc.add_paragraph('La formule de calcul utilisée est la suivante:')
                        formula = r'n = \frac{Z^2 \cdot p \cdot (1 - p)}{e^2}'
                        img_buffer = generate_latex_image(formula)
                        doc.add_picture(img_buffer, width=Inches(3))
                        doc.add_paragraph("Avec n la taille de l'échantillon, Z le niveau de confiance, p la proportion estimée et e la marge d'erreur")

                        doc.add_heading('Méthodologie', 2)
                        doc.add_paragraph("La méthode utilisée dans le calcul de la taille de l'échantillon est la méthode de Cochran")

                    # doc.add_picture('path_to_graph.png', width=Inches(5))
                    # Sauvegarder le fichier Word en mémoire
                    report_io = io.BytesIO()
                    doc.save(report_io)
                    report_io.seek(0)
                    return report_io
                st.markdown("### Rapport des résultats")

                # Télécharger le rapport si l'utilisateur clique sur le bouton
                if st.session_state.type == 'Stratifie':
                    report_file = generate_stratified_report()
                # Télécharger le rapport
                st.download_button(label="Télécharger le Rapport", data=report_file, file_name="rapport_sondage.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            with colo:
                def to_excel(df):
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                    return output.getvalue()

                if st.session_state.base_echant is not None:
                    st.markdown('### Télécharger la base de données')

                    if st.session_state.type == 'Simple':
                        excel_data = to_excel(st.session_state.base_echant)

                # Ajouter un bouton de téléchargement
                    st.download_button(
                        label="Télécharger l'échantillon sous format Excel",
                        data=excel_data,
                        file_name='Echantillon.xlsx',
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    )
                elif st.session_state.base_echant_strat is not None:
                    if st.session_state.type == 'Stratifie':
                        excel_data = to_excel(st.session_state.base_echant_strat)

                # Ajouter un bouton de téléchargement
                    st.download_button(
                        label="Télécharger l'échantillon sous format Excel",
                        data=excel_data,
                        file_name='Echantillon.xlsx',
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    )









