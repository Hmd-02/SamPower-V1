import streamlit as st
import pandas as pd
from fpdf import FPDF
#import sympy as sp
import matplotlib.pyplot as plt
import io
import os
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime


# Vérifier si l'utilisateur a accès à cette page
st.title("Résultats")
if 'marge' not in st.session_state:
    st.session_state.marge = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse' not in st.session_state:
    st.session_state.tauxReponse = None
if 'tailleEffet' not in st.session_state:
    st.session_state.tailleEffet = None
if 'parametre' not in st.session_state:
    st.session_state.parametre = None
if 'methode' not in st.session_state:
    st.session_state.methode = None
if 'confiance' not in st.session_state:
    st.session_state.confiance = None
if 'sigma' not in st.session_state:
    st.session_state.sigma = None
if 'prop_est' not in st.session_state:
    st.session_state.prop_est = None
if 'echantillon' not in st.session_state:
    st.session_state.echantillon = None
if 'calcul' not in st.session_state:
    st.session_state.calcul = False
if 'base_echant' not in st.session_state:
    st.session_state.base_echant = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population' not in st.session_state:
    st.session_state.population = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'type' not in st.session_state:
    st.session_state.type = None
if 'marge_strat' not in st.session_state:
    st.session_state.marge_strat = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse_strat' not in st.session_state:
    st.session_state.tauxReponse_strat = None
if 'tailleEffet_strat' not in st.session_state:
    st.session_state.tailleEffet_strat = None
if 'parametre_strat' not in st.session_state:
    st.session_state.parametre_strat = None
if 'methode_strat' not in st.session_state:
    st.session_state.methode_strat = None
if 'confiance_strat' not in st.session_state:
    st.session_state.confiance_strat = None
if 'sigma_strat' not in st.session_state:
    st.session_state.sigma_strat = None
if 'prop_est_strat' not in st.session_state:
    st.session_state.prop_est_strat = None
if 'calcul_strat' not in st.session_state:
    st.session_state.calcul_strat = False
if 'echantillon_strat' not in st.session_state:
    st.session_state.echantillon_strat = None
if 'base_echant_strat' not in st.session_state:
    st.session_state.base_echant_strat = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population_strat' not in st.session_state:
    st.session_state.population_strat = None
if 'variable_strat' not in st.session_state:
    st.session_state.variable_strat = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'modality_count' not in st.session_state:
    st.session_state.modality_count = None
if 'resultat' not in st.session_state:
    st.session_state.resultat = None
if 'tauxSondage' not in st.session_state:
    st.session_state.tauxSondage = None
if 'tauxSondage_strat' not in st.session_state:
    st.session_state.tauxSondage_strat = None
if 'tot_taille' not in st.session_state:
    st.session_state.tot_taille = None



# Fonction pour générer une image à partir d'une formule LaTeX
def generate_latex_image(formula, fontsize=30):
    fig, ax = plt.subplots()
    ax.text(0.5, 0.5, f'${formula}$', fontsize=fontsize, ha='center', va='center')
    ax.axis('off')

    # Sauvegarder l'image dans un buffer mémoire
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer

# Fonction pour générer le rapport dans un fichier Word
if st.session_state.type == 'Simple':
    def generate_report():
        # Créer un document Word
        doc = Document()
        # Ajouter un titre
        doc.add_heading('Rapport des résultats', 0)
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        doc.add_paragraph(f'Généré le : {current_time}')
        # Texte introductif
        doc.add_paragraph("L'application SamPower est un outil puissant dédié à la réalisation de sondages et à l'estimation des paramètres statistiques essentiels. Conçue pour être à la fois intuitive et accessible, elle permet aux utilisateurs d'effectuer des échantillonnages de manière rapide et efficace, que ce soit pour des sondages aléatoires simples ou stratifiés. Son importance réside dans sa capacité à offrir des résultats fiables tout en simplifiant le processus complexe de calcul de taille d'échantillon, un élément clé pour garantir des résultats statistiquement représentatifs dans le cadre d'études quantitatives. Le présent rapport est structuré de la manière suivante :")
        doc.add_paragraph("- Présentation des paramètres d'entrée et des résultats : Cette section regroupe l'ensemble des résultats obtenus ainsi que les différents paramètres entrés par l'utilisateur.")            
        doc.add_paragraph("- Présentation des explications et de la méthodologie utilisée : Ici, nous détaillons la méthodologie employée pour générer les résultats, notamment la logique derrière les formules et les méthodes statistiques appliquées.")
        # Ajouter des informations sur la méthode
        doc.add_heading("Présentation des paramètres d'entrée et des résultats", level = 1)
        doc.add_heading('Paramètre à estimer :', level=2)
        doc.add_paragraph(f"Paramètre : {st.session_state.parametre}")
        # Ajouter les paramètres d'entrée
        doc.add_heading('Paramètres d\'entrée :', level=2)
        if st.session_state.parametre == "Proportion":
            doc.add_paragraph(f"Proportion estimée : {st.session_state.prop_est}")
        if st.session_state.parametre == "Moyenne":
            doc.add_paragraph(f"Ecart-type estimé : {st.session_state.sigma}")
        doc.add_paragraph(f"Erreur Marginale : {st.session_state.marge}")
        doc.add_paragraph(f"Niveau de confiance : {st.session_state.confiance * 100}%")
        doc.add_paragraph(f"Taille de la population : {st.session_state.population}")
        doc.add_paragraph(f"Taux de réponse : {st.session_state.tauxReponse * 100}%")
        doc.add_paragraph(f"Effet de plan : {st.session_state.tailleEffet * 100}%")
        if st.session_state.population is not None:
            doc.add_paragraph(f"Taux de sondage : {st.session_state.tauxSondage * 100}%")


        # Afficher les résultats d'entrée
        doc.add_heading('Résultats :', level=2)
        doc.add_paragraph(f"Taille de l'échantillon calculée : {st.session_state.echantillon}")
        doc.add_page_break()

        # Ajouter des explications
        doc.add_heading('Explications et méthodologie utilisée', level=1)
        doc.add_heading('Explication des résultats :', level=2)
        doc.add_paragraph(f"Sur la base des paramètres définis par l'utilisateur, tels que la marge d'erreur, le niveau de confiance et la taille de la population, le calcul a déterminé qu'un échantillon optimal de {st.session_state.echantillon} individus est nécessaire pour garantir que les résultats de l'enquête soient représentatifs. Ce nombre permet de s'assurer que les estimations issues de l'enquête respectent la précision statistique souhaitée tout en maintenant un niveau de confiance élevé. Ce calcul prend également en compte les hypothèses formulées sur la proportion attendue dans la population.")

    # Formule de la méthode de Wald
        if st.session_state.parametre == 'Proportion':
            doc.add_heading('Formule de calcul', level = 2)
            doc.add_paragraph('La formule de calcul utilisée est la suivante:')
            formula = r'n = \frac{Z^2 \cdot p \cdot (1 - p)}{e^2}'
            img_buffer = generate_latex_image(formula)
            doc.add_picture(img_buffer, width=Inches(3))
            doc.add_paragraph("Avec n la taille de l'échantillon, Z le niveau de confiance, p la proportion estimée et e la marge d'erreur")
            
            doc.add_heading('Méthodologie', level = 2)
            doc.add_paragraph("La méthode utilisée dans le calcul de la taille de l'échantillon est la méthode de Wald")

        # Sauvegarder le fichier Word en mémoire
        report_io = io.BytesIO()
        doc.save(report_io)
        report_io.seek(0)
        return report_io


# Fonction pour générer le rapport dans un fichier Word
if st.session_state.type == 'Stratifie': 
    def generate_stratified_report():
        # Créer un document Word
        doc = Document()
        # Ajouter un titre
        doc.add_heading('Rapport des résultats (Sondage Stratifié)', 0)
        
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        doc.add_paragraph(f'Généré le : {current_time}')   
        doc.add_paragraph("L'application SamPower est un outil puissant dédié à la réalisation de sondages et à l'estimation des paramètres statistiques essentiels. Conçue pour être à la fois intuitive et accessible, elle permet aux utilisateurs d'effectuer des échantillonnages de manière rapide et efficace, que ce soit pour des sondages aléatoires simples ou stratifiés. Son importance réside dans sa capacité à offrir des résultats fiables tout en simplifiant le processus complexe de calcul de taille d'échantillon, un élément clé pour garantir des résultats statistiquement représentatifs dans le cadre d'études quantitatives. Le présent rapport est structuré de la manière suivante :")
        doc.add_paragraph("- Présentation des paramètres d'entrée et des résultats : Cette section regroupe l'ensemble des résultats obtenus ainsi que les différents paramètres entrés par l'utilisateur.")            
        doc.add_paragraph("- Présentation des explications et de la méthodologie utilisée : Ici, nous détaillons la méthodologie employée pour générer les résultats, notamment la logique derrière les formules et les méthodes statistiques appliquées.")
        
        # Ajouter des informations sur la méthode
        doc.add_heading("Présentation des paramètres d'entrée et des résultats", level = 1)
        doc.add_heading('Paramètres à estimer :', level=2)
        doc.add_paragraph(f"Paramètre : {st.session_state.parametre_strat}")

        # Ajouter les paramètres d'entrée
        doc.add_heading('Paramètres d\'entrée :', level=2)
        doc.add_paragraph(f"Niveau de confiance : {st.session_state.confiance_strat * 100}%")
        doc.add_paragraph(f"Erreur Marginale : {st.session_state.marge_strat}")
        doc.add_paragraph(f"Taux de réponse : {st.session_state.tauxReponse_strat * 100}%")
        if st.session_state.type_alloc == 'Allocation proportionnelle':
            doc.add_paragraph(f"Taux de sondage : {st.session_state.tauxSondage_strat * 100}%")

    
        if st.session_state.type_alloc == 'Calcul par strate':
            if isinstance(st.session_state.modality_count, pd.DataFrame):
                # Ajouter les paramètres sous forme de tableau dans le document Word
                doc.add_heading('Taille et effet de plan :', level=3)
                table = doc.add_table(rows=1, cols=len(st.session_state.modality_count.columns))
                hdr_cells = table.rows[0].cells
                for i, column in enumerate(st.session_state.modality_count.columns):
                    hdr_cells[i].text = column

                for index, row in st.session_state.modality_count.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)


        # Afficher les résultats d'entrée
        doc.add_heading('Résultats :', level=2)
        if st.session_state.type_alloc == 'Allocation proportionnelle':
            doc.add_paragraph(f"Taille de l'échantillon calculée : {st.session_state.echantillon_strat}")
        if st.session_state.type_alloc == 'Calcul par strate':
            result_df = st.session_state.resultat
        if st.session_state.type_alloc == 'Allocation proportionnelle':
            result_df = st.session_state.modality_count
        doc.add_page_break()
        if isinstance(result_df, pd.DataFrame):
            # Ajouter les paramètres sous forme de tableau dans le document Word
            doc.add_heading('Tableau des résultats', level=3)
            table = doc.add_table(rows=1, cols=len(result_df.columns))
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(result_df.columns):
                hdr_cells[i].text = column

            for index, row in result_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        # Ajouter des explications
        
        doc.add_heading('Explications et méthodologie utilisée', level=1)
        doc.add_heading('Explication des résultats :', level=2)
        if st.session_state.type_alloc == 'Allocation proportionnelle':
            doc.add_paragraph(f"Suite à l'application de la méthode de sondage stratifié, l'échantillon total a été fixé à {st.session_state.echantillon_strat}. Cet échantillon a été réparti de manière proportionnelle entre les différentes strates en fonction de leur taille dans la population de référence comme présenté dans le tableau des résultats.Cette répartition garantit que chaque strate est représentée de manière adéquate dans l'échantillon final, permettant ainsi des estimations précises et fiables pour chacune des strates tout en minimisant les biais potentiels liés à la taille des sous-populations. En conséquence, les résultats obtenus à partir de cet échantillon seront statistiquement représentatifs de l'ensemble de la population cible tout en respectant la structure spécifique des strates.")            
        if st.session_state.type_alloc == 'Calcul par strate':
            doc.add_paragraph(f"Dans le cadre de ce sondage stratifié, la taille de l'échantillon total a été directement déterminée au sein de chaque strate selon des critères spécifiques à chacune d'elles, tels que leur proportion ou leur variabilité interne. L'échantillon est donc de {st.session_state.tot_taille}. Cette méthode permet de s'assurer que chaque strate est représentée de manière proportionnée à ses caractéristiques propres, garantissant ainsi une précision accrue des estimations dans chaque sous-groupe de la population. En optant pour un calcul direct des tailles d'échantillon dans chaque strate, nous maximisons la représentativité de l'échantillon pour chaque groupe cible, tout en prenant en compte les particularités de chaque sous-population. Ce processus offre des résultats plus détaillés et pertinents pour chaque strate, assurant ainsi une meilleure compréhension des dynamiques internes de la population étudiée. ")
        
        # Formule de la méthode de Wald
        if st.session_state.parametre_strat == 'Proportion':
            doc.add_heading('Formule de calcul',2)
            doc.add_paragraph('La formule de calcul utilisée est la suivante:')
            formula = r'n = \frac{Z^2 \cdot p \cdot (1 - p)}{e^2}'
            img_buffer = generate_latex_image(formula)
            doc.add_picture(img_buffer, width=Inches(3))
            doc.add_paragraph("Avec n la taille de l'échantillon, Z le niveau de confiance, p la proportion estimée et e la marge d'erreur")

            doc.add_heading('Méthodologie', 2)
            doc.add_paragraph("La méthode utilisée dans le calcul de la taille de l'échantillon est la méthode de Cochran")

        
        
        # doc.add_picture('path_to_graph.png', width=Inches(5))
        
        # Sauvegarder le fichier Word en mémoire
        report_io = io.BytesIO()
        doc.save(report_io)
        report_io.seek(0)
        return report_io

# Interface Streamlit
st.markdown("### Rapport des résultats")

# Télécharger le rapport si l'utilisateur clique sur le bouton
if st.session_state.type == 'Simple':
    report_file = generate_report()
if st.session_state.type == 'Stratifie':
    report_file = generate_stratified_report()

# Télécharger le rapport
st.download_button(label="Télécharger le Rapport", data=report_file, file_name="rapport_sondage.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# Fonction pour créer un fichier Excel en mémoire
def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    return output.getvalue()

if st.session_state.base_echant is not None:
    st.markdown('### Télécharger la base de données')

    if st.session_state.type == 'Simple':
        excel_data = to_excel(st.session_state.base_echant)

# Ajouter un bouton de téléchargement
    st.download_button(
        label="Télécharger l'échantillon sous format Excel",
        data=excel_data,
        file_name='Echantillon.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
elif st.session_state.base_echant_strat is not None:
    if st.session_state.type == 'Stratifie':
        excel_data = to_excel(st.session_state.base_echant_strat)

# Ajouter un bouton de téléchargement
    st.download_button(
        label="Télécharger l'échantillon sous format Excel",
        data=excel_data,
        file_name='Echantillon.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )







