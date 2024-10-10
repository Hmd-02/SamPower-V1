import streamlit as st
import pandas as pd
#import numpy as np
import samplics

from samplics.sampling import SampleSize
from samplics.utils.types import SizeMethod, PopParam
from st_aggrid import AgGrid
import io
from fpdf import FPDF
#import sympy as sp
import matplotlib.pyplot as plt
import io
import os
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime



# Vérifier si l'utilisateur a accès à cette page
st.markdown("<h1 style='color:#2fd5cf; text-align: center;'>Sondage aléatoire simple </h1>", unsafe_allow_html=True)
if 'marge' not in st.session_state:
    st.session_state.marge = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse' not in st.session_state:
    st.session_state.tauxReponse = None
if 'tailleEffet' not in st.session_state:
    st.session_state.tailleEffet = None
if 'parametre' not in st.session_state:
    st.session_state.parametre = None
if 'methode' not in st.session_state:
    st.session_state.methode = None
if 'confiance' not in st.session_state:
    st.session_state.confiance = None
if 'sigma' not in st.session_state:
    st.session_state.sigma = None
if 'prop_est' not in st.session_state:
    st.session_state.prop_est = None
if 'calcul_moy' not in st.session_state:
    st.session_state.calcul = False
if 'calcul_prop' not in st.session_state:
    st.session_state.calcul= False
if 'base_echant_moy' not in st.session_state:
    st.session_state.base_echant = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population' not in st.session_state:
    st.session_state.population = None
if 'tauxSondage' not in st.session_state:
    st.session_state.tauxSondage = None
#if 'base_telecharger' not in st.session_state:
#    st.session_state.base_telecharger = None

if 'marge' not in st.session_state:
    st.session_state.marge = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse' not in st.session_state:
    st.session_state.tauxReponse = None
if 'tailleEffet' not in st.session_state:
    st.session_state.tailleEffet = None
if 'parametre' not in st.session_state:
    st.session_state.parametre = None
if 'methode' not in st.session_state:
    st.session_state.methode = None
if 'confiance' not in st.session_state:
    st.session_state.confiance = None
if 'sigma' not in st.session_state:
    st.session_state.sigma = None
if 'prop_est' not in st.session_state:
    st.session_state.prop_est = None
if 'echantillon' not in st.session_state:
    st.session_state.echantillon = None
if 'calcul' not in st.session_state:
    st.session_state.calcul = False
if 'base_echant' not in st.session_state:
    st.session_state.base_echant = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population' not in st.session_state:
    st.session_state.population = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'type' not in st.session_state:
    st.session_state.type = None
if 'marge_strat' not in st.session_state:
    st.session_state.marge_strat = None
if 'taillePop' not in st.session_state:
    st.session_state.taillePop = None
if 'tauxReponse_strat' not in st.session_state:
    st.session_state.tauxReponse_strat = None
if 'tailleEffet_strat' not in st.session_state:
    st.session_state.tailleEffet_strat = None
if 'parametre_strat' not in st.session_state:
    st.session_state.parametre_strat = None
if 'methode_strat' not in st.session_state:
    st.session_state.methode_strat = None
if 'confiance_strat' not in st.session_state:
    st.session_state.confiance_strat = None
if 'sigma_strat' not in st.session_state:
    st.session_state.sigma_strat = None
if 'prop_est_strat' not in st.session_state:
    st.session_state.prop_est_strat = None
if 'calcul_strat' not in st.session_state:
    st.session_state.calcul_strat = False
if 'echantillon_strat' not in st.session_state:
    st.session_state.echantillon_strat = None
if 'base_echant_strat' not in st.session_state:
    st.session_state.base_echant_strat = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'population_strat' not in st.session_state:
    st.session_state.population_strat = None
if 'variable_strat' not in st.session_state:
    st.session_state.variable_strat = None
if 'type_alloc' not in st.session_state:
    st.session_state.type_alloc = None
if 'modality_count' not in st.session_state:
    st.session_state.modality_count = None
if 'resultat' not in st.session_state:
    st.session_state.resultat = None
if 'tauxSondage' not in st.session_state:
    st.session_state.tauxSondage = None
if 'tauxSondage_strat' not in st.session_state:
    st.session_state.tauxSondage_strat = None
if 'tot_taille' not in st.session_state:
    st.session_state.tot_taille = None

st.info("Les champs avec l'astérisque (*) rouge doivent être obligatoirement remplis")
st.markdown("##### Paramètre à estimer <span style='color:red;'>*</span>", unsafe_allow_html=True)
st.session_state.parametre = st.selectbox('Paramètre à estimer', ['None', 'Proportion', 'Moyenne'], label_visibility='collapsed')
if st.session_state.parametre == 'Proportion' or st.session_state.parametre == 'Moyenne':
    col1, col2,col3 = st.columns(3)

    with col1:
        if st.session_state.parametre == "Moyenne":
            st.markdown("##### Écart-type estimé <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Si vous avez accès à des données provenant d'études antérieures sur la même population ou des populations similaires, vous pouvez utiliser l'écart-type observé dans ces études comme estimation de 𝜎σ.")
            st.session_state.sigma = st.number_input("Marge d'erreur ", min_value = 0, label_visibility='collapsed')

        if st.session_state.parametre == "Proportion":
            st.markdown("##### Proportion estimée <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("Si vous disposez d'études similaires précédentes vous pouvez utiliser ces résultats. Sinon la valeur prise par défaut est 0.5")
            st.session_state.prop_est = st.number_input("Proportion estimée", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')

        if st.session_state.parametre == "Proportion":
            st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
            st.session_state.marge = st.number_input("Marge d'erreur", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')
        
        if st.session_state.parametre == "Moyenne":
            st.markdown("##### Marge d'erreur <span style='color:red;'>*</span>", unsafe_allow_html=True)
            with st.expander("En savoir plus ..."):
                st.info("La marge d'erreur permet de savoir correspond a la moitié de l'étendue de l'intervalle de confiance. C'est la valeur qui est ajoutée et retranchée de la proportion estimée pour obtenir l'intervalle de confiance")
            st.session_state.marge = st.number_input("Marge d'erreur moy ", min_value= 0, label_visibility='collapsed')
    
    with col2:
        st.markdown("##### Niveau de confiance <span style='color:red;'>*</span>", unsafe_allow_html=True)
        with st.expander("En savoir plus ..."):
            st.info("Le niveau de confiance représente la probabilité que vous êtes prêt à accepter quant à la possibilité que votre estimation soit incorrecte.")
        st.session_state.confiance = st.number_input("Niveau de confiance",min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')
        
        st.markdown("##### Taille de la population")
        with st.expander("En savoir plus ..."):
            st.info('Taille de la population : Le nombre total d’individus dans la population étudiée. Ce paramètre est utilisé pour calculer la taille de l’échantillon.')
        st.session_state.population = st.number_input("Taille de la population", min_value = 1, value=st.session_state.taillePop, label_visibility='collapsed')

    
    with col3:
        st.markdown("##### Taux de réponse <span style='color:red;'>*</span>", unsafe_allow_html=True)
        with st.expander("En savoir plus ..."):
            st.info("Le taux de réponse est le pourcentage de personnes ayant répondu à un sondage par rapport au nombre total de personnes sollicitées. Assurez-vous d'atteindre ce taux de réponse au risque de ne pas avoir des estimations précises.")
        st.session_state.tauxReponse = st.number_input("Taux de réponse", min_value=0.0, max_value= 1.0, step = 0.0001, format="%.4f", label_visibility='collapsed')

        st.markdown("##### Effet de plan <span style='color:red;'>*</span>", unsafe_allow_html=True)
        with st.expander("En savoir plus ..."):
            st.info("L'effet de plan est un facteur utilisé dans les enquêtes pour ajuster la taille d'échantillon nécessaire lorsque le plan d'échantillonnage diffère d'un simple échantillonnage aléatoire. Il mesure l'impact du plan d'échantillonnage sur la précision des estimations. Si vous utilisez un plan d'échantillonnage stratifié ou en grappes, par exemple, l'effet de plan prend en compte la corrélation au sein des groupes ou des strates, ce qui peut augmenter la variance des estimations. Un effet de plan supérieur à 1 signifie que la taille de l'échantillon doit être augmentée pour maintenir la même précision. Comme nous sommes dans le cas d'un sondage aléatoire simple, l'effet de plan est de 1")
        st.session_state.tailleEffet =  st.selectbox("Effet de plan",[1], label_visibility='collapsed')




    # Calcul de la taille dans le cas d'une proportion
    if st.session_state.parametre == "Proportion":
        st.session_state.methode = SampleSize(param=PopParam.prop, method=SizeMethod.wald, strat=False)
    if st.session_state.parametre == "Proportion":
        if st.session_state.methode is not None and st.session_state.tailleEffet != 0 and st.session_state.tauxReponse != 0 and st.session_state.taillePop != 0 and st.session_state.marge != 0:
            def calcul_taille():
                st.session_state.methode.calculate(
                    target = st.session_state.prop_est, 
                    half_ci=st.session_state.marge,
                    deff=st.session_state.tailleEffet, 
                    resp_rate=st.session_state.tauxReponse,
                    pop_size = st.session_state.population,
                    alpha = 1 - st.session_state.confiance)
                st.session_state.echantillon = st.session_state.methode.samp_size
            st.session_state.calcul = st.button("Calculer la taille",help="Cliquer pour lancer le calcul de la taille",on_click=calcul_taille)


    # Calcul de la taille dans le cas d'une moyenne
    if st.session_state.parametre == "Moyenne":
        st.session_state.methode = SampleSize(param=PopParam.mean, method=SizeMethod.wald, strat=False)
    if st.session_state.parametre == "Moyenne":
        if st.session_state.methode is not None and st.session_state.tailleEffet != 0 and st.session_state.tauxReponse != 0 and st.session_state.taillePop != 0 and st.session_state.marge != 0:
            def calcul_taille():
                st.session_state.methode.calculate(
                    sigma = st.session_state.sigma,
                    half_ci=st.session_state.marge,
                    deff=st.session_state.tailleEffet, 
                    resp_rate=st.session_state.tauxReponse,
                    pop_size = st.session_state.population,
                    alpha = 1 - st.session_state.confiance)
                st.session_state.echantillon = st.session_state.methode.samp_size
            st.session_state.calcul = st.button("Calculer la taille",help="Cliquer pour lancer le calcul de la taille",on_click=calcul_taille)

    if st.session_state.calcul:
        st.markdown("<h1 style='color:#2fd5cf; text-align: center;'>Résultats </h1>", unsafe_allow_html=True)
        cola, colb, colc= st.columns(3)
        with cola:
            st.write(f"#### **Taille d'échantillon**:", st.session_state.echantillon)
            if st.session_state.population is not None:
                st.session_state.tauxSondage = round(st.session_state.echantillon / st.session_state.population,4)
                st.write("**Taux de sondage**: ",st.session_state.tauxSondage)
            if st.session_state.df is not None:
                with st.expander("Plus d'options"):
                    st.write("*Echantillon tiré*",unsafe_allow_html=True)
                    st.session_state.base_echant = st.session_state.df.sample(n = st.session_state.echantillon, random_state = 42)
                    AgGrid(st.session_state.base_echant)
            
        with colb:
            # Fonction pour générer une image à partir d'une formule LaTeX
            def generate_latex_image(formula, fontsize=30):
                fig, ax = plt.subplots()
                ax.text(0.5, 0.5, f'${formula}$', fontsize=fontsize, ha='center', va='center')
                ax.axis('off')

                # Sauvegarder l'image dans un buffer mémoire
                img_buffer = io.BytesIO()
                plt.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
                plt.close(fig)
                img_buffer.seek(0)
                
                return img_buffer

            # Fonction pour générer le rapport dans un fichier Word
            if st.session_state.type == 'Simple':
                def generate_report():
                    # Créer un document Word
                    doc = Document()
                    # Ajouter un titre
                    doc.add_heading('Rapport des résultats', 0)
                    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    doc.add_paragraph(f'Généré le : {current_time}')
                    # Texte introductif
                    doc.add_paragraph("L'application SamPower est un outil puissant dédié à la réalisation de sondages et à l'estimation des paramètres statistiques essentiels. Conçue pour être à la fois intuitive et accessible, elle permet aux utilisateurs d'effectuer des échantillonnages de manière rapide et efficace, que ce soit pour des sondages aléatoires simples ou stratifiés. Son importance réside dans sa capacité à offrir des résultats fiables tout en simplifiant le processus complexe de calcul de taille d'échantillon, un élément clé pour garantir des résultats statistiquement représentatifs dans le cadre d'études quantitatives. Le présent rapport est structuré de la manière suivante :")
                    doc.add_paragraph("- Présentation des paramètres d'entrée et des résultats : Cette section regroupe l'ensemble des résultats obtenus ainsi que les différents paramètres entrés par l'utilisateur.")            
                    doc.add_paragraph("- Présentation des explications et de la méthodologie utilisée : Ici, nous détaillons la méthodologie employée pour générer les résultats, notamment la logique derrière les formules et les méthodes statistiques appliquées.")
                    # Ajouter des informations sur la méthode
                    doc.add_heading("Présentation des paramètres d'entrée et des résultats", level = 1)
                    doc.add_heading('Paramètre à estimer :', level=2)
                    doc.add_paragraph(f"Paramètre : {st.session_state.parametre}")
                    # Ajouter les paramètres d'entrée
                    doc.add_heading('Paramètres d\'entrée :', level=2)
                    if st.session_state.parametre == "Proportion":
                        doc.add_paragraph(f"Proportion estimée : {st.session_state.prop_est}")
                    if st.session_state.parametre == "Moyenne":
                        doc.add_paragraph(f"Ecart-type estimé : {st.session_state.sigma}")
                    doc.add_paragraph(f"Erreur Marginale : {st.session_state.marge}")
                    doc.add_paragraph(f"Niveau de confiance : {st.session_state.confiance * 100}%")
                    doc.add_paragraph(f"Taille de la population : {st.session_state.population}")
                    doc.add_paragraph(f"Taux de réponse : {st.session_state.tauxReponse * 100}%")
                    doc.add_paragraph(f"Effet de plan : {st.session_state.tailleEffet * 100}%")
                    if st.session_state.population is not None:
                        doc.add_paragraph(f"Taux de sondage : {st.session_state.tauxSondage * 100}%")


                    # Afficher les résultats d'entrée
                    doc.add_heading('Résultats :', level=2)
                    doc.add_paragraph(f"Taille de l'échantillon calculée : {st.session_state.echantillon}")
                    doc.add_page_break()

                    # Ajouter des explications
                    doc.add_heading('Explications et méthodologie utilisée', level=1)
                    doc.add_heading('Explication des résultats :', level=2)
                    doc.add_paragraph(f"Sur la base des paramètres définis par l'utilisateur, tels que la marge d'erreur, le niveau de confiance et la taille de la population, le calcul a déterminé qu'un échantillon optimal de {st.session_state.echantillon} individus est nécessaire pour garantir que les résultats de l'enquête soient représentatifs. Ce nombre permet de s'assurer que les estimations issues de l'enquête respectent la précision statistique souhaitée tout en maintenant un niveau de confiance élevé. Ce calcul prend également en compte les hypothèses formulées sur la proportion attendue dans la population.")

                # Formule de la méthode de Wald
                    if st.session_state.parametre == 'Proportion':
                        doc.add_heading('Formule de calcul', level = 2)
                        doc.add_paragraph('La formule de calcul utilisée est la suivante:')
                        formula = r'n = \frac{Z^2 \cdot p \cdot (1 - p)}{e^2}'
                        img_buffer = generate_latex_image(formula)
                        doc.add_picture(img_buffer, width=Inches(3))
                        doc.add_paragraph("Avec n la taille de l'échantillon, Z le niveau de confiance, p la proportion estimée et e la marge d'erreur")
                        
                        doc.add_heading('Méthodologie', level = 2)
                        doc.add_paragraph("La méthode utilisée dans le calcul de la taille de l'échantillon est la méthode de Wald")

                    # Sauvegarder le fichier Word en mémoire
                    report_io = io.BytesIO()
                    doc.save(report_io)
                    report_io.seek(0)
                    return report_io
            #st.markdown("### Rapport des résultats")

            # Télécharger le rapport si l'utilisateur clique sur le bouton
            if st.session_state.type == 'Simple':
                report_file = generate_report()
            # Télécharger le rapport
            st.download_button(label="Télécharger le Rapport", data=report_file, file_name="rapport_sondage.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with colc:
            # Fonction pour créer un fichier Excel en mémoire
            def to_excel(df):
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Sheet1')
                return output.getvalue()

            if st.session_state.base_echant is not None:
                #st.markdown('### Télécharger la base de données')

                if st.session_state.type == 'Simple':
                    excel_data = to_excel(st.session_state.base_echant)

            # Ajouter un bouton de téléchargement
                st.download_button(
                    label="Télécharger l'échantillon sous format Excel",
                    data=excel_data,
                    file_name='Echantillon.xlsx',
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )





                        









